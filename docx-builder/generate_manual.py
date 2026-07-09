# pyrefly: ignore [missing-import]
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def main():
    doc = docx.Document()

    # Set up normal font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # Title Section
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TimeForge Workforce Workspace")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5) # Indigo 600

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("System Administration & Multi-Role User Manual")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b) # Slate 500

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Section 1: System Startup
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Getting the System Running")
    h1_run.font.bold = True
    h1_run.font.size = Pt(18)
    h1_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) # Slate 900
    
    p = doc.add_paragraph("TimeForge is built as a decoupled web application with a Laravel API backend, React frontend SPA, PostgreSQL database, Redis caching, and ClamAV upload scanner. Follow these instructions to initialize and run the system locally.")
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading("Prerequisites", level=2)
    p = doc.add_paragraph()
    p.add_run("- Docker & Docker Compose: ").bold = True
    p.add_run("Used to run database, redis, and malware scanner containers.\n")
    p.add_run("- Node.js (v20+): ").bold = True
    p.add_run("Required to build and run the frontend React server locally.\n")
    p.add_run("- PHP 8.2+: ").bold = True
    p.add_run("Required if executing CLI scripts directly on the host.")

    doc.add_heading("Startup Steps", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Step 1: Start Docker Containers\n").bold = True
    p.add_run("Run this command in the root project directory to start the PostgreSQL database, Redis store, and ClamAV daemon:\n")
    p.add_run("docker compose up -d").italic = True
    p.paragraph_format.left_indent = Inches(0.25)
    
    p = doc.add_paragraph()
    p.add_run("Step 2: Initialize Backend Database\n").bold = True
    p.add_run("Run the database migrations and seed default demo accounts:\n")
    p.add_run("docker compose exec app php artisan migrate\n").italic = True
    p.add_run("docker compose exec app php artisan db:seed").italic = True
    p.paragraph_format.left_indent = Inches(0.25)

    p = doc.add_paragraph()
    p.add_run("Step 3: Launch Frontend Application\n").bold = True
    p.add_run("Open a new terminal session, navigate to the frontend directory, install dependencies, and start the development server:\n")
    p.add_run("cd frontend && npm install && npm run dev").italic = True
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph("The frontend will now be accessible at http://localhost:5173, and the API gateway will serve requests at http://localhost:8000.")

    # Table of Seeded Demo Credentials
    doc.add_heading("Default Demo Credentials", level=2)
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Role", "Username (Email)", "Password"]
    for i, title_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = title_text
        set_cell_background(cell, "4F46E5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    rows_data = [
        ["Administrator", "admin@timeforge.test", "Passw0rd123!"],
        ["Supervisor", "supervisor@timeforge.test", "Passw0rd123!"],
        ["HR & Finance", "hr@timeforge.test", "Passw0rd123!"],
        ["Employee", "employee@timeforge.test", "Passw0rd123!"]
    ]

    for row_idx, row_data in enumerate(rows_data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = text
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            else:
                set_cell_background(cell, "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Section 2: Role-based Guides
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("2. Role-Based User Manual")
    h1_run.font.bold = True
    h1_run.font.size = Pt(18)
    h1_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    doc.add_heading("A. Employee Role", level=2)
    p = doc.add_paragraph("As an Employee, you use TimeForge to capture your daily productivity metrics, log timesheets, and review KPI targets.")
    
    p = doc.add_paragraph()
    p.add_run("1. Punch In/Out: ").bold = True
    p.add_run("Use the header clock widget to start your daily shifts. The system will record your exact start and stop times.\n")
    p.add_run("2. Create Time Entries: ").bold = True
    p.add_run("Log specific tasks you performed during the day under Time Tracking. You can link your entries to clients and projects.\n")
    p.add_run("3. Submit Timesheets: ").bold = True
    p.add_run("At the end of your billing cycle, navigate to Time Tracking, click 'Submit Timesheet', and specify the date. This locks your entries and notifies your supervisor for approval.\n")
    p.add_run("4. Daily Scrum updates: ").bold = True
    p.add_run("Input your previous work, planned work, blockers, and notes to share visibility with your team.\n")
    p.add_run("5. Leave Requests: ").bold = True
    p.add_run("Request time off by specifying date ranges and reasons under My Leave.")
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("B. Supervisor Role", level=2)
    p = doc.add_paragraph("Supervisors manage specific departments, review submitted team timesheets, and coordinate blocker resolutions.")
    
    p = doc.add_paragraph()
    p.add_run("1. Approve/Reject Timesheets: ").bold = True
    p.add_run("Under Team Timesheets, view timesheets submitted by your department. You can Approve, Reject (requires a comment), or request a Revision (unlocks entries to let the employee make updates).\n")
    p.add_run("2. Track Scrum Updates: ").bold = True
    p.add_run("Monitor daily updates from your team. You can leave comments on updates to assist team members with blockers.\n")
    p.add_run("3. Manage Team Leave: ").bold = True
    p.add_run("Review pending vacation and time-off requests from department staff under Team Leave.\n")
    p.add_run("4. View Department KPIs: ").bold = True
    p.add_run("Inspect current team assignments and completion trajectories.")
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("C. HR / Finance Role", level=2)
    p = doc.add_paragraph("HR and Finance users monitor labor expenditures, validate organization payroll readiness, and export compiled records.")
    
    p = doc.add_paragraph()
    p.add_run("1. Monitor Payroll Costs: ").bold = True
    p.add_run("View overall regular and overtime hours, alongside estimated payroll costs for active employees.\n")
    p.add_run("2. Review Exceptions: ").bold = True
    p.add_run("Track employees with unsubmitted days, open timers, missing hourly rates, or excessive overtime under Payroll Exceptions.\n")
    p.add_run("3. Export Reports: ").bold = True
    p.add_run("Download payroll summary worksheets as PDF or Excel spreadsheets. Exports are compiled asynchronously in background queues; you will receive an in-app notification once the download link is generated.")
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("D. Administrator Role", level=2)
    p = doc.add_paragraph("Administrators hold full access to system parameters, directories, audit records, and security features.")
    
    p = doc.add_paragraph()
    p.add_run("1. Approve Registrations: ").bold = True
    p.add_run("Review and verify incoming registration requests under Account Approvals.\n")
    p.add_run("2. Manage Users & Roles: ").bold = True
    p.add_run("Add new profiles, update role permissions, modify department assignments, and toggle status configurations (Active/Deactivated).\n" )
    p.add_run("3. Global Directories: ").bold = True
    p.add_run("Set up system departments, clients, projects, global holidays, and KPI structures.\n")
    p.add_run("4. Inspect Audit Logs: ").bold = True
    p.add_run("Track every database edit, profile update, timesheet submission, and system access attempt for compliance audits.\n")
    p.add_run("5. Company Settings: ").bold = True
    p.add_run("Modify company settings, configure custom branding logos, and configure global rules.")
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Section 3: Enterprise Security Features
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("3. Enterprise Security & Privacy Controls")
    h1_run.font.bold = True
    h1_run.font.size = Pt(18)
    h1_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    p = doc.add_paragraph("TimeForge includes strict security controls to safeguard proprietary business details:")
    p = doc.add_paragraph()
    p.add_run("- Malware Scanning: ").bold = True
    p.add_run("All uploaded files (avatars, timesheet attachments) are streamed to a ClamAV scanner socket before being committed to storage. Viruses are automatically blocked and return a warning.\n")
    p.add_run("- Two-Factor Authentication: ").bold = True
    p.add_run("Administrators and HR/Finance users are protected by secondary login checks. Upon password verification, an OTP code is emailed and required to authenticate.\n")
    p.add_run("- AI Data Masking: ").bold = True
    p.add_run("To ensure employee privacy, all source records are recursively scrubbed and anonymized (e.g. replacing real names/emails with 'Employee A', 'Employee B') before leaving our servers to hit Google's Gemini AI API gateway.")
    p.paragraph_format.left_indent = Inches(0.25)

    # Save Manual
    doc.save("TimeForge_User_Manual.docx")
    print("TimeForge_User_Manual.docx generated successfully!")

if __name__ == '__main__':
    main()
